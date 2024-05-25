import datetime
import re
from docx2pdf import convert
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement, ns

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = 2 # Right Aligned

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

def read_html_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    return html_content

def create_word_document(html_content, output_path):
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    header_footer_data = "SECRET"

    # Create a new section and set margins
    section = doc.sections[0]
    section.left_margin = section.right_margin = int(1440 * 200)
    section.top_margin = section.bottom_margin = int(1440 * 200)

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{header_footer_data}"
    header_para.alignment = 2 # Right Aligned

    # Find all HTML tags in the HTML content
    html_tags = soup.find_all('html')
    for idx, html_tag in enumerate(html_tags, start=1):
        if idx > 1:
            doc.add_page_break()
        add_html_content(html_tag, doc)

    apply_styles_from_css(doc, soup)

    # FOOTER
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"{header_footer_data}\t"
    footer_para.alignment = 1
    add_page_number(footer.paragraphs[0])

    # Save the document
    doc.save(output_path)

def apply_styles_from_css(doc, soup):
    # Find and parse the style tag
    style_tag = soup.find('style')
    if style_tag:
        css_rules = style_tag.string.split("}")
        for rule in css_rules:
            rule_parts = rule.strip().split("{")
            if len(rule_parts) == 2:
                selector = rule_parts[0].strip()
                properties = rule_parts[1].strip().split(";")
                for prop in properties:
                    prop_parts = prop.strip().split(":")
                    if len(prop_parts) == 2:
                        apply_style(doc, selector, prop_parts[0].strip(), prop_parts[1].strip())

def apply_style(doc, selector, prop_name, prop_value):
    if selector == 'table':
        pass
    elif selector.startswith('.'):
        # Apply style to paragraphs with class
        class_name = selector[1:]
        for paragraph in doc.paragraphs:
            if paragraph.style.name == class_name:
                if prop_name == 'text-align':
                    if prop_value == 'center':
                        paragraph.alignment = 1  # CENTER
                    elif prop_value == 'right':
                        paragraph.alignment = 2  # RIGHT
    elif re.match(r'h\d+', selector):
        # Apply style to headers
        level = int(selector[1:])
        style_name = 'Heading {}'.format(level)
        for paragraph in doc.paragraphs:
            if paragraph.style.name == style_name:
                if prop_name == 'text-align':
                    if prop_value == 'center':
                        paragraph.alignment = 1  # CENTER
                    elif prop_value == 'right':
                        paragraph.alignment = 2  # RIGHT
    elif selector == 'th':
        # Apply style to table headers
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.paragraphs[0].alignment = 1  # CENTER

def add_html_content(html_tag, doc):
    # Process the current tag
    process_html_tag(html_tag, doc)

    # Recursively process the children of the current tag
    for child in html_tag.children:
        if isinstance(child, Tag):
            add_html_content(child, doc)

def process_html_tag(tag, doc):
    if re.match(r"^h\d+$", tag.name):
        # Handle header tags
        level = int(tag.name[1])
        add_header_with_format(tag, doc, level)
    elif tag.name == 'div':
        # Handle div tags
        handle_div(tag, doc)
    elif tag.name == 'p':
        # Handle paragraph tags
        add_paragraph_with_format(tag, doc)
    elif tag.name == 'table':
        # Handle table tags
        handle_table(tag, doc)
        doc.add_paragraph("\n")
    elif tag.name == 'img':
        # Handle Image Tag
        img_src = tag['src']
        width = int(tag['width']) if 'width' in tag.attrs else None
        height = int(tag['height']) if 'height' in tag.attrs else None

        add_image_to_docx(doc, img_src, width=width, height=height)

def add_image_to_docx(doc, img_path, width=None, height=None):
    if width and height:
        doc.add_picture(img_path, width=Inches(width / 96), height=Inches(height / 96))  # Convert pixels to inches
    else:
        doc.add_picture(img_path)

def handle_div(div_tag, doc):
    pass


def add_paragraph_with_format(element, doc):
    paragraph = doc.add_paragraph()
    for item in element.contents:
        if isinstance(item, str):
            # If the item is a string, add it as text
            paragraph.add_run(item)
        elif item.name == 'b':
            # If the item is a <b> tag, add its text as bold
            run = paragraph.add_run(item.text)
            run.bold = True
        elif item.name == 'br':
            # If the item is a <br> tag, add a line break
            paragraph.add_run().add_break()

def add_header_with_format(element, doc, level):
    style = 'Heading {}'.format(level)
    paragraph = doc.add_paragraph(style=style)
    for item in element.contents:
        if isinstance(item, str):
            paragraph.add_run(item)
        elif item.name == 'b':
            run = paragraph.add_run(item.text)
            run.bold = True

def handle_table(table_tag, doc):
    # Create a new table in the document
    rows = table_tag.find_all('tr')
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    table = doc.add_table(rows=0, cols=max_cols)
    table.style = 'Table Grid'  # Set table style with borders

    # Add table header row
    th_tags = table_tag.find_all('th')
    header_row = table.add_row().cells
    for i, th_tag in enumerate(th_tags):
        header_cell = header_row[i]
        header_cell.text = th_tag.text.strip()
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_bg_color(header_cell, RGBColor(242, 242, 242))

    # Add table body rows
    for tr_tag in table_tag.find_all('tr'):
        row_cells = [cell.text.strip().replace(",", ",\n") for cell in tr_tag.find_all('td')]
        if any(row_cells):  # Check if the row has any content
            new_row = table.add_row().cells
            for i, cell_text in enumerate(row_cells):
                new_row[i].text = cell_text

def set_cell_bg_color(cell, rgb_color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{color}"/>'.format(nsdecls('w'), color=rgb_color))
    cell._element.tcPr.append(shading_elm)

def docx_to_pdf(docx_file, pdf_file):
    convert(docx_file, pdf_file)

if __name__ == "__main__":
    html_file_path = "" # html filename
    output_docx_path = "" # doc filename
    output_pdf_path = "" # pdf filename

    start_time = datetime.datetime.now()

    print("DOCUMENT PROCESSING....")

    html_content = read_html_file(html_file_path)
    create_word_document(html_content, output_docx_path)

    end_time = datetime.datetime.now()

    print(f"DOCUMENT CREATED!!!!!!\n{(end_time-start_time).total_seconds()} seconds took in html to doc conversion")

    start_time = datetime.datetime.now()

    docx_to_pdf(output_docx_path,output_pdf_path)

    end_time = datetime.datetime.now()

    print(f"PDF CREATED!!!!!!\n{(end_time-start_time).total_seconds()} seconds took in doc to pdf conversion")